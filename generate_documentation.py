#!/usr/bin/env python3
"""
Script to generate ZenFlow app documentation in DOCX format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style hyperlink
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def create_zenflow_documentation():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    title = doc.add_heading('ZENFLOW', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.bold = True
    
    subtitle = doc.add_paragraph('Pomodoro Timer Application')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project Information
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Mobile Application Development Project\n').font.size = Pt(14)
    info.add_run('\n')
    info.add_run('Built with Flutter & Clean Architecture\n').font.size = Pt(12)
    info.add_run('\n\n')
    info.add_run('Version 1.0.0\n').font.size = Pt(12)
    info.add_run(f'October 2025').font.size = Pt(12)
    
    doc.add_page_break()
    
    # Table of Contents
    toc = doc.add_heading('TABLE OF CONTENTS', 1)
    toc.runs[0].font.size = Pt(14)
    toc.runs[0].font.bold = True
    
    toc_items = [
        ('1.', 'INTRODUCTION'),
        ('   1.1', 'Project Overview'),
        ('   1.2', 'Purpose and Scope'),
        ('   1.3', 'Target Users'),
        ('2.', 'SYSTEM FEATURES'),
        ('   2.1', 'Pomodoro Timer'),
        ('   2.2', 'Task Management'),
        ('   2.3', 'Cycle Tracking'),
        ('   2.4', 'Settings Configuration'),
        ('   2.5', 'Audio Notifications'),
        ('3.', 'TECHNICAL ARCHITECTURE'),
        ('   3.1', 'Clean Architecture'),
        ('   3.2', 'MVVM Pattern'),
        ('   3.3', 'Project Structure'),
        ('   3.4', 'Technology Stack'),
        ('4.', 'USER INTERFACE'),
        ('   4.1', 'Design Principles'),
        ('   4.2', 'Main Screen'),
        ('   4.3', 'Timer Controls'),
        ('   4.4', 'Task List'),
        ('5.', 'DATA MANAGEMENT'),
        ('   5.1', 'Local Storage'),
        ('   5.2', 'Data Models'),
        ('   5.3', 'Persistence Strategy'),
        ('6.', 'IMPLEMENTATION DETAILS'),
        ('   6.1', 'Domain Layer'),
        ('   6.2', 'Data Layer'),
        ('   6.3', 'Presentation Layer'),
        ('   6.4', 'Core Components'),
        ('7.', 'USER GUIDE'),
        ('   7.1', 'Getting Started'),
        ('   7.2', 'Using the Timer'),
        ('   7.3', 'Managing Tasks'),
        ('   7.4', 'Customizing Settings'),
        ('8.', 'INSTALLATION AND SETUP'),
        ('   8.1', 'Prerequisites'),
        ('   8.2', 'Installation Steps'),
        ('   8.3', 'Running the Application'),
        ('9.', 'TESTING'),
        ('   9.1', 'Testing Strategy'),
        ('   9.2', 'Widget Tests'),
        ('10.', 'FUTURE ENHANCEMENTS'),
        ('11.', 'CONCLUSION'),
        ('12.', 'REFERENCES'),
    ]
    
    for num, item in toc_items:
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'{num} {item}')
    
    doc.add_page_break()
    
    # 1. INTRODUCTION
    doc.add_heading('1. INTRODUCTION', 1)
    
    doc.add_heading('1.1 Project Overview', 2)
    p = doc.add_paragraph()
    p.add_run(
        'ZenFlow is a modern Pomodoro Timer application developed using Flutter framework with '
        'Clean Architecture principles. The application helps users manage their time effectively '
        'using the Pomodoro Technique, a time management method that uses a timer to break work '
        'into intervals, traditionally 25 minutes in length, separated by short breaks.'
    )
    
    p = doc.add_paragraph()
    p.add_run(
        'The application is built with maintainability, scalability, and testability in mind, '
        'implementing Clean Architecture patterns and the MVVM (Model-View-ViewModel) design pattern. '
        'It uses Hive, a lightweight and fast NoSQL database, for local data persistence, ensuring '
        'that user data is preserved across app sessions without requiring internet connectivity.'
    )
    
    doc.add_heading('1.2 Purpose and Scope', 2)
    p = doc.add_paragraph()
    p.add_run('The primary objectives of ZenFlow are:')
    
    purposes = [
        'Provide an intuitive and distraction-free Pomodoro timer interface',
        'Enable users to track multiple tasks and their associated Pomodoro sessions',
        'Visualize productivity through cycle tracking and progress indicators',
        'Offer customizable timer durations to suit individual work patterns',
        'Store all user data locally for privacy and offline accessibility',
        'Deliver audio notifications for timer completion events',
        'Maintain a clean and aesthetically pleasing user interface'
    ]
    
    for purpose in purposes:
        doc.add_paragraph(purpose, style='List Bullet')
    
    doc.add_heading('1.3 Target Users', 2)
    p = doc.add_paragraph()
    p.add_run(
        'ZenFlow is designed for individuals who want to improve their productivity and time management, including:'
    )
    
    users = [
        'Students working on academic projects and assignments',
        'Professionals managing multiple work tasks',
        'Freelancers tracking billable hours',
        'Remote workers maintaining focus during work-from-home sessions',
        'Anyone looking to adopt the Pomodoro Technique for better time management'
    ]
    
    for user in users:
        doc.add_paragraph(user, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. SYSTEM FEATURES
    doc.add_heading('2. SYSTEM FEATURES', 1)
    
    doc.add_heading('2.1 Pomodoro Timer', 2)
    p = doc.add_paragraph()
    p.add_run(
        'The core feature of ZenFlow is its Pomodoro timer functionality, which implements '
        'the traditional Pomodoro Technique with three distinct timer modes:'
    )
    
    timer_modes = [
        ('Focus Session', '25 minutes of concentrated work time', 'Primary work period for maximum productivity'),
        ('Short Break', '5 minutes of rest between focus sessions', 'Brief respite to maintain energy levels'),
        ('Long Break', '15 minutes of extended rest', 'Earned after completing 4 focus sessions (1 complete cycle)')
    ]
    
    for mode, duration, description in timer_modes:
        p = doc.add_paragraph()
        p.add_run(f'{mode}: ').bold = True
        p.add_run(f'{duration} - {description}')
    
    p = doc.add_paragraph()
    p.add_run(
        'The timer displays a large, easily readable countdown with MM:SS format. It features '
        'intuitive START/PAUSE/STOP controls and automatically transitions between timer types '
        'based on the Pomodoro cycle logic.'
    )
    
    doc.add_heading('2.2 Task Management', 2)
    p = doc.add_paragraph()
    p.add_run('ZenFlow includes a comprehensive task management system that allows users to:')
    
    task_features = [
        'Create new tasks with descriptive titles',
        'Mark tasks as complete with a single tap',
        'Edit existing task titles',
        'Delete tasks that are no longer needed',
        'Track the number of Pomodoros completed for each task',
        'View separate lists for active and completed tasks',
        'Associate timer sessions with specific tasks'
    ]
    
    for feature in task_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run(
        'Tasks are displayed in a scrollable list with visual indicators for completion status '
        'and Pomodoro counts. The task list integrates seamlessly with the timer, allowing users '
        'to select which task they are currently working on.'
    )
    
    doc.add_heading('2.3 Cycle Tracking', 2)
    p = doc.add_paragraph()
    p.add_run(
        'The application automatically tracks Pomodoro cycles, displaying visual progress indicators:'
    )
    
    cycle_features = [
        'Current cycle progress (1-4 focus sessions)',
        'Completed cycles counter',
        'Total focus sessions completed',
        'Visual dots showing cycle status:',
        '  • Green checkmarks for completed focus sessions',
        '  • Red pulsing ring for current focus session',
        '  • Grey dots for upcoming focus sessions'
    ]
    
    for feature in cycle_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run(
        'After completing 4 focus sessions (1 complete cycle), the timer automatically suggests '
        'a long break, helping users maintain sustainable productivity patterns.'
    )
    
    doc.add_heading('2.4 Settings Configuration', 2)
    p = doc.add_paragraph()
    p.add_run('Users can customize their Pomodoro experience through the settings panel:')
    
    settings_options = [
        'Focus Duration: Adjustable work session length (default 25 minutes)',
        'Short Break Duration: Customizable short break length (default 5 minutes)',
        'Long Break Duration: Configurable long break length (default 15 minutes)',
        'Cycles Before Long Break: Number of focus sessions before long break (default 4)',
        'Auto Switch: Automatically start next timer when current one completes',
        'Reset All: Option to reset all cycle counters and start fresh'
    ]
    
    for option in settings_options:
        doc.add_paragraph(option, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run(
        'All settings are persisted locally and apply immediately without requiring app restart.'
    )
    
    doc.add_heading('2.5 Audio Notifications', 2)
    p = doc.add_paragraph()
    p.add_run(
        'ZenFlow provides audio feedback when timer sessions complete, alerting users even if '
        'the app is in the background. The notification system includes:'
    )
    
    audio_features = [
        'System sound playback when timer reaches zero',
        'Modal dialog displaying completion message',
        'Context-aware messaging (e.g., "Focus session complete! Time for a short break")',
        'Option to dismiss or immediately start the next timer',
        'Support for auto-start with visual confirmation'
    ]
    
    for feature in audio_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. TECHNICAL ARCHITECTURE
    doc.add_heading('3. TECHNICAL ARCHITECTURE', 1)
    
    doc.add_heading('3.1 Clean Architecture', 2)
    p = doc.add_paragraph()
    p.add_run(
        'ZenFlow implements Clean Architecture, a software design philosophy that separates the '
        'application into layers with distinct responsibilities. This architecture ensures:'
    )
    
    clean_arch_benefits = [
        'Independence from frameworks and UI',
        'Testability of business logic',
        'Independence from external agencies',
        'Maintainable and scalable codebase',
        'Clear separation of concerns'
    ]
    
    for benefit in clean_arch_benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('The application is structured into three main layers:').bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Domain Layer:').bold = True
    p = doc.add_paragraph()
    p.add_run(
        'Contains the core business logic and entities. This layer is completely independent of '
        'Flutter and external libraries, consisting of pure Dart code. It defines:'
    )
    
    domain_components = [
        'Entities: Task, TimerSession, AppSettings',
        'Repository interfaces: Abstract contracts for data operations',
        'Business rules and domain logic'
    ]
    
    for component in domain_components:
        doc.add_paragraph(component, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Data Layer:').bold = True
    p = doc.add_paragraph()
    p.add_run(
        'Implements the repository interfaces defined in the domain layer. It handles all data '
        'persistence operations using Hive database:'
    )
    
    data_components = [
        'Models: Hive-annotated classes with type adapters for serialization',
        'Repository implementations: Concrete implementations of domain repositories',
        'Database operations: CRUD operations for tasks, sessions, and settings'
    ]
    
    for component in data_components:
        doc.add_paragraph(component, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Presentation Layer:').bold = True
    p = doc.add_paragraph()
    p.add_run(
        'Contains all UI-related code and follows the MVVM pattern:'
    )
    
    presentation_components = [
        'ViewModels: Manage UI state using ChangeNotifier',
        'Screens: Complete app pages (e.g., HomeScreen)',
        'Widgets: Reusable UI components',
        'State management: Provider package for dependency injection and state propagation'
    ]
    
    for component in presentation_components:
        doc.add_paragraph(component, style='List Bullet')
    
    doc.add_heading('3.2 MVVM Pattern', 2)
    p = doc.add_paragraph()
    p.add_run(
        'The Model-View-ViewModel (MVVM) pattern separates the UI from business logic:'
    )
    
    mvvm_components = [
        ('Model', 'Domain entities and data models representing app state'),
        ('View', 'Flutter widgets that display data and capture user interactions'),
        ('ViewModel', 'Intermediary layer exposing data streams and handling user actions')
    ]
    
    for component, description in mvvm_components:
        p = doc.add_paragraph()
        p.add_run(f'{component}: ').bold = True
        p.add_run(description)
    
    p = doc.add_paragraph()
    p.add_run('ZenFlow implements three main ViewModels:')
    
    viewmodels = [
        'TimerViewModel: Manages timer state, countdown, and cycle logic',
        'TaskViewModel: Handles task CRUD operations and state',
        'SettingsViewModel: Manages app configuration and preferences'
    ]
    
    for vm in viewmodels:
        doc.add_paragraph(vm, style='List Bullet')
    
    doc.add_heading('3.3 Project Structure', 2)
    p = doc.add_paragraph()
    p.add_run('The codebase is organized into a logical folder structure:')
    
    structure = """
lib/
├── core/
│   ├── constants/
│   │   └── app_constants.dart
│   ├── services/
│   │   └── audio_service.dart
│   └── theme/
│       ├── app_colors.dart
│       └── app_theme.dart
├── data/
│   ├── models/
│   │   ├── task_model.dart
│   │   ├── timer_session_model.dart
│   │   └── settings_model.dart
│   └── repositories/
│       ├── task_repository_impl.dart
│       ├── session_repository_impl.dart
│       └── settings_repository_impl.dart
├── domain/
│   ├── entities/
│   │   ├── task.dart
│   │   ├── timer_session.dart
│   │   └── app_settings.dart
│   └── repositories/
│       ├── task_repository.dart
│       ├── session_repository.dart
│       └── settings_repository.dart
├── presentation/
│   ├── screens/
│   │   └── home_screen.dart
│   ├── viewmodels/
│   │   ├── timer_viewmodel.dart
│   │   ├── task_viewmodel.dart
│   │   └── settings_viewmodel.dart
│   └── widgets/
│       ├── timer_display.dart
│       ├── timer_controls.dart
│       ├── timer_type_selector.dart
│       ├── cycle_progress.dart
│       ├── task_list_widget.dart
│       ├── settings_widget.dart
│       ├── timer_complete_dialog.dart
│       └── footer_widget.dart
└── main.dart
"""
    
    p = doc.add_paragraph(structure, style='No Spacing')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)
    
    doc.add_heading('3.4 Technology Stack', 2)
    
    tech_stack = [
        ('Flutter', '3.5.3+', 'Cross-platform UI framework'),
        ('Dart', 'Latest', 'Programming language'),
        ('Hive', '2.2.3', 'Lightweight NoSQL database'),
        ('Provider', '6.1.1', 'State management solution'),
        ('Equatable', '2.0.5', 'Value equality for entities'),
        ('Google Fonts', '6.1.0', 'Custom typography (Poppins font)'),
        ('Audioplayers', '5.2.1', 'Audio playback for notifications'),
        ('Build Runner', '2.4.7', 'Code generation tool'),
        ('Hive Generator', '2.0.1', 'Hive adapter code generation')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technology'
    hdr_cells[1].text = 'Version'
    hdr_cells[2].text = 'Purpose'
    
    for tech, version, purpose in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = version
        row_cells[2].text = purpose
    
    doc.add_page_break()
    
    # 4. USER INTERFACE
    doc.add_heading('4. USER INTERFACE', 1)
    
    doc.add_heading('4.1 Design Principles', 2)
    p = doc.add_paragraph()
    p.add_run('ZenFlow\'s UI design follows modern mobile app design principles:')
    
    design_principles = [
        'Dark theme to reduce eye strain during extended use',
        'Minimalist interface to minimize distractions',
        'Large, readable typography for easy visibility',
        'Color-coded timer types (red for focus, yellow for short break, blue for long break)',
        'Intuitive controls with clear visual feedback',
        'Responsive layout that adapts to different screen sizes',
        'Consistent spacing and alignment throughout the app'
    ]
    
    for principle in design_principles:
        doc.add_paragraph(principle, style='List Bullet')
    
    doc.add_heading('4.2 Main Screen', 2)
    p = doc.add_paragraph()
    p.add_run('The home screen is divided into several key sections:')
    
    screen_sections = [
        ('Header', 'App title "ZenFlow" with fire emoji icon and "Pomodoro Timer" subtitle'),
        ('Timer Display', 'Large circular timer showing countdown in MM:SS format with color-coded underline'),
        ('Timer Type Selector', 'Three tabs for Focus/Short Break/Long Break with smooth sliding animation'),
        ('Timer Controls', 'START/PAUSE and RESET buttons with clear labeling'),
        ('Cycle Progress', 'Visual dots showing progress through current cycle (1-4 focus sessions)'),
        ('Statistics', 'Display of completed cycles and total focus sessions'),
        ('Task List', 'Scrollable list of active and completed tasks'),
        ('Footer', 'Developer credit and app information')
    ]
    
    for section, description in screen_sections:
        p = doc.add_paragraph()
        p.add_run(f'{section}: ').bold = True
        p.add_run(description)
    
    doc.add_heading('4.3 Timer Controls', 2)
    p = doc.add_paragraph()
    p.add_run('The timer control interface provides:')
    
    control_features = [
        'START button: Initiates the countdown timer',
        'PAUSE button: Temporarily halts the timer (replaces START when running)',
        'RESET button: Resets timer to initial duration',
        'Visual feedback: Button states change based on timer status',
        'Disabled states: Controls are disabled appropriately to prevent invalid actions'
    ]
    
    for feature in control_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('4.4 Task List', 2)
    p = doc.add_paragraph()
    p.add_run('The task management interface features:')
    
    task_ui_features = [
        'Input field at the top for adding new tasks',
        'Separate sections for active and completed tasks',
        'Checkbox to mark tasks complete/incomplete',
        'Pomodoro counter badge showing completed sessions per task',
        'Edit icon to modify task titles',
        'Delete icon to remove tasks',
        'Smooth animations for task state changes'
    ]
    
    for feature in task_ui_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 5. DATA MANAGEMENT
    doc.add_heading('5. DATA MANAGEMENT', 1)
    
    doc.add_heading('5.1 Local Storage', 2)
    p = doc.add_paragraph()
    p.add_run(
        'ZenFlow uses Hive as its local database solution. Hive is a lightweight, fast NoSQL '
        'database written in pure Dart, making it ideal for Flutter applications. Key advantages:'
    )
    
    hive_benefits = [
        'No native dependencies required',
        'Works on all platforms (mobile, desktop, web)',
        'Fast read/write operations',
        'Type-safe with code generation',
        'Minimal boilerplate code',
        'Encrypted storage support',
        'Lazy box loading for improved performance'
    ]
    
    for benefit in hive_benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_heading('5.2 Data Models', 2)
    p = doc.add_paragraph()
    p.add_run('The application persists three main data models:')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('TaskModel:').bold = True
    
    task_fields = [
        'id: String - Unique identifier',
        'title: String - Task description',
        'completed: bool - Completion status',
        'createdAt: DateTime - Creation timestamp',
        'pomodorosCompleted: int - Number of completed Pomodoros'
    ]
    
    for field in task_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('TimerSessionModel:').bold = True
    
    session_fields = [
        'id: String - Unique identifier',
        'timerTypeIndex: int - Type (0=focus, 1=short, 2=long)',
        'durationMinutes: int - Session duration',
        'startTime: DateTime - Session start time',
        'endTime: DateTime? - Session end time (nullable)',
        'completed: bool - Whether session was completed',
        'taskId: String? - Associated task ID (nullable)'
    ]
    
    for field in session_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('SettingsModel:').bold = True
    
    settings_fields = [
        'id: String - Unique identifier',
        'focusDuration: int - Focus session minutes (default 25)',
        'shortBreakDuration: int - Short break minutes (default 5)',
        'longBreakDuration: int - Long break minutes (default 15)',
        'cyclesBeforeLongBreak: int - Cycles before long break (default 4)',
        'autoSwitch: bool - Auto-start next timer (default false)'
    ]
    
    for field in settings_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('5.3 Persistence Strategy', 2)
    p = doc.add_paragraph()
    p.add_run('Data persistence is handled through repository implementations:')
    
    persistence_strategy = [
        'Repositories use Hive boxes for data storage',
        'Each model type has its own dedicated box',
        'CRUD operations are asynchronous to prevent UI blocking',
        'Data is automatically persisted on device storage',
        'ViewModels interact only with repository interfaces, not direct Hive access',
        'Settings are loaded on app startup and cached in memory',
        'Task and session data is loaded on-demand'
    ]
    
    for strategy in persistence_strategy:
        doc.add_paragraph(strategy, style='List Bullet')
    
    doc.add_page_break()
    
    # 6. IMPLEMENTATION DETAILS
    doc.add_heading('6. IMPLEMENTATION DETAILS', 1)
    
    doc.add_heading('6.1 Domain Layer', 2)
    p = doc.add_paragraph()
    p.add_run(
        'The domain layer contains pure business logic with no external dependencies. '
        'Entities use Equatable for value comparison:'
    )
    
    entity_example = """
class Task extends Equatable {
  final String id;
  final String title;
  final bool completed;
  final DateTime createdAt;
  final int pomodorosCompleted;

  const Task({
    required this.id,
    required this.title,
    required this.completed,
    required this.createdAt,
    this.pomodorosCompleted = 0,
  });

  @override
  List<Object?> get props => [id, title, completed, createdAt, pomodorosCompleted];
}
"""
    
    p = doc.add_paragraph(entity_example, style='No Spacing')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)
    
    doc.add_heading('6.2 Data Layer', 2)
    p = doc.add_paragraph()
    p.add_run(
        'Data models extend HiveObject and use type adapters generated by build_runner. '
        'Repository implementations provide concrete data operations:'
    )
    
    repo_example = """
class TaskRepositoryImpl implements TaskRepository {
  final Box<TaskModel> _box;

  TaskRepositoryImpl(this._box);

  @override
  Future<List<Task>> getTasks() async {
    return _box.values.map((model) => model.toEntity()).toList();
  }

  @override
  Future<void> addTask(Task task) async {
    final model = TaskModel.fromEntity(task);
    await _box.put(task.id, model);
  }
  
  // Additional CRUD methods...
}
"""
    
    p = doc.add_paragraph(repo_example, style='No Spacing')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)
    
    doc.add_heading('6.3 Presentation Layer', 2)
    p = doc.add_paragraph()
    p.add_run(
        'ViewModels extend ChangeNotifier to provide reactive state management. '
        'The TimerViewModel handles the core timer logic:'
    )
    
    viewmodel_features = [
        'Countdown timer using Dart Timer with 1-second intervals',
        'Automatic cycle progression logic',
        'Integration with audio service for notifications',
        'Callback system for UI alerts',
        'State preservation during pause/resume',
        'Settings integration for customizable durations'
    ]
    
    for feature in viewmodel_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('6.4 Core Components', 2)
    p = doc.add_paragraph()
    p.add_run('AudioService:').bold = True
    p.add_run(
        ' Singleton service handling audio playback using the audioplayers package. '
        'Plays system alert sounds when timers complete.'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('AppConstants:').bold = True
    p.add_run(
        ' Central location for app-wide constants including box names, default durations, '
        'and configuration values.'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('AppColors:').bold = True
    p.add_run(
        ' Defines the color palette used throughout the app, ensuring visual consistency. '
        'Includes colors for focus (red), short break (yellow), long break (blue), and UI elements.'
    )
    
    doc.add_page_break()
    
    # 7. USER GUIDE
    doc.add_heading('7. USER GUIDE', 1)
    
    doc.add_heading('7.1 Getting Started', 2)
    p = doc.add_paragraph()
    p.add_run('Upon first launch, ZenFlow initializes with default settings:')
    
    defaults = [
        'Focus duration: 25 minutes',
        'Short break: 5 minutes',
        'Long break: 15 minutes',
        'Cycles before long break: 4',
        'Timer starts in Focus mode',
        'No tasks created by default'
    ]
    
    for default in defaults:
        doc.add_paragraph(default, style='List Bullet')
    
    doc.add_heading('7.2 Using the Timer', 2)
    
    timer_steps = [
        ('Select Timer Type', 'Tap Focus, Short Break, or Long Break tab at the top'),
        ('Start Timer', 'Press the START button to begin countdown'),
        ('Pause if Needed', 'Tap PAUSE to temporarily stop (STOP appears when paused)'),
        ('Reset Timer', 'Use RESET button to return to initial duration'),
        ('Wait for Completion', 'Timer will alert you with sound and dialog when complete'),
        ('Continue Cycle', 'Choose to start next timer or dismiss and take a break')
    ]
    
    for step, description in timer_steps:
        p = doc.add_paragraph()
        p.add_run(f'Step: {step}').bold = True
        p.add_run(f' - {description}')
    
    doc.add_heading('7.3 Managing Tasks', 2)
    
    task_steps = [
        ('Create Task', 'Enter task title in input field and tap add button'),
        ('Mark Complete', 'Tap checkbox next to task to toggle completion'),
        ('Edit Task', 'Tap edit icon (pencil) to modify task title'),
        ('Delete Task', 'Tap delete icon (trash) to remove task'),
        ('View Pomodoros', 'See Pomodoro count badge on each task'),
        ('Associate with Timer', 'Tasks automatically track when timer is running')
    ]
    
    for step, description in task_steps:
        p = doc.add_paragraph()
        p.add_run(f'{step}: ').bold = True
        p.add_run(description)
    
    doc.add_heading('7.4 Customizing Settings', 2)
    
    settings_steps = [
        ('Open Settings', 'Tap floating settings button (gear icon) in bottom-right'),
        ('Adjust Durations', 'Use number inputs to change timer lengths'),
        ('Set Cycle Count', 'Change number of focus sessions before long break'),
        ('Toggle Auto-Switch', 'Enable/disable automatic timer progression'),
        ('Apply Changes', 'Tap APPLY button to save settings'),
        ('Reset All', 'Use RESET ALL to clear all cycles and start fresh')
    ]
    
    for step, description in settings_steps:
        p = doc.add_paragraph()
        p.add_run(f'{step}: ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # 8. INSTALLATION AND SETUP
    doc.add_heading('8. INSTALLATION AND SETUP', 1)
    
    doc.add_heading('8.1 Prerequisites', 2)
    p = doc.add_paragraph()
    p.add_run('Before installing ZenFlow, ensure you have:')
    
    prerequisites = [
        'Flutter SDK (version 3.5.3 or higher)',
        'Dart SDK (included with Flutter)',
        'Android Studio or VS Code with Flutter extensions',
        'Git for version control',
        'Android SDK for Android development',
        'Xcode for iOS development (macOS only)'
    ]
    
    for prereq in prerequisites:
        doc.add_paragraph(prereq, style='List Bullet')
    
    doc.add_heading('8.2 Installation Steps', 2)
    
    install_steps = [
        'Clone the repository or download source code',
        'Navigate to project directory in terminal',
        'Run "flutter pub get" to install dependencies',
        'Run "flutter pub run build_runner build --delete-conflicting-outputs" to generate Hive adapters',
        'Connect device or start emulator',
        'Run "flutter run" to launch the app'
    ]
    
    for i, step in enumerate(install_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_heading('8.3 Running the Application', 2)
    p = doc.add_paragraph()
    p.add_run('Platform-specific run commands:')
    
    run_commands = [
        ('Android', 'flutter run -d <device-id>'),
        ('iOS', 'flutter run -d <device-id>'),
        ('Web', 'flutter run -d chrome'),
        ('Windows', 'flutter run -d windows'),
        ('macOS', 'flutter run -d macos'),
        ('Linux', 'flutter run -d linux')
    ]
    
    for platform, command in run_commands:
        p = doc.add_paragraph()
        p.add_run(f'{platform}: ').bold = True
        p.add_run(command)
        p.runs[-1].font.name = 'Courier New'
    
    doc.add_page_break()
    
    # 9. TESTING
    doc.add_heading('9. TESTING', 1)
    
    doc.add_heading('9.1 Testing Strategy', 2)
    p = doc.add_paragraph()
    p.add_run(
        'ZenFlow implements automated testing to ensure code quality and reliability. '
        'The testing strategy includes:'
    )
    
    testing_types = [
        'Unit Tests: Test individual functions and classes in isolation',
        'Widget Tests: Test UI components and user interactions',
        'Integration Tests: Test complete user flows',
        'Manual Testing: Test on real devices for UX validation'
    ]
    
    for test_type in testing_types:
        doc.add_paragraph(test_type, style='List Bullet')
    
    doc.add_heading('9.2 Widget Tests', 2)
    p = doc.add_paragraph()
    p.add_run(
        'Widget tests verify that UI components render correctly and respond to user input. '
        'The test suite includes:'
    )
    
    widget_tests = [
        'Smoke test verifying app launches successfully',
        'Timer display rendering with correct format',
        'Button state changes based on timer status',
        'Task list operations (add, edit, delete, complete)',
        'Settings dialog functionality',
        'Navigation between screens'
    ]
    
    for test in widget_tests:
        doc.add_paragraph(test, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Run tests with: ')
    p.add_run('flutter test').font.name = 'Courier New'
    
    doc.add_page_break()
    
    # 10. FUTURE ENHANCEMENTS
    doc.add_heading('10. FUTURE ENHANCEMENTS', 1)
    p = doc.add_paragraph()
    p.add_run(
        'While ZenFlow is fully functional, several enhancements are planned for future versions:'
    )
    
    enhancements = [
        ('Statistics Dashboard', 'Detailed analytics showing productivity trends over time, including charts and graphs'),
        ('Data Export', 'Export task and session data to CSV or JSON for external analysis'),
        ('Cloud Sync', 'Optional cloud backup and sync across multiple devices'),
        ('Multiple Themes', 'Additional color schemes including light theme option'),
        ('Push Notifications', 'Background notifications when app is not active'),
        ('Custom Timer Presets', 'Save and quickly switch between different timer configurations'),
        ('Task Categories', 'Organize tasks into projects or categories'),
        ('Desktop Widget', 'System tray widget for Windows/macOS/Linux'),
        ('Pomodoro History', 'Detailed log of all completed sessions'),
        ('Focus Mode', 'Distraction blocking during focus sessions'),
        ('Integration with Calendar', 'Sync scheduled tasks with calendar apps'),
        ('Team Features', 'Shared tasks and collaborative Pomodoro sessions')
    ]
    
    for title, description in enhancements:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # 11. CONCLUSION
    doc.add_heading('11. CONCLUSION', 1)
    p = doc.add_paragraph()
    p.add_run(
        'ZenFlow represents a modern approach to implementing the Pomodoro Technique in a '
        'mobile application. By leveraging Flutter\'s cross-platform capabilities and Clean '
        'Architecture principles, the application delivers a robust, maintainable, and user-friendly '
        'productivity tool.'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        'The application successfully achieves its core objectives:'
    )
    
    achievements = [
        'Provides intuitive timer interface following Pomodoro Technique best practices',
        'Enables comprehensive task tracking integrated with timer sessions',
        'Offers customizable settings to accommodate individual work patterns',
        'Maintains clean architecture for long-term maintainability',
        'Ensures data persistence without requiring internet connectivity',
        'Delivers smooth, responsive user experience across platforms'
    ]
    
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        'ZenFlow demonstrates how modern software development practices—including Clean Architecture, '
        'MVVM pattern, and comprehensive testing—can be applied to create applications that are both '
        'technically sound and user-focused. The modular structure ensures that new features can be '
        'added without disrupting existing functionality, while the separation of concerns makes the '
        'codebase accessible to developers of varying experience levels.'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        'As productivity tools continue to evolve, ZenFlow provides a solid foundation that can grow '
        'with user needs while maintaining the simplicity and focus that makes the Pomodoro Technique '
        'effective.'
    )
    
    doc.add_page_break()
    
    # 12. REFERENCES
    doc.add_heading('12. REFERENCES', 1)
    
    references = [
        ('Flutter Documentation', 'https://flutter.dev/docs'),
        ('Dart Language Tour', 'https://dart.dev/guides/language/language-tour'),
        ('Hive Database', 'https://docs.hivedb.dev/'),
        ('Provider Package', 'https://pub.dev/packages/provider'),
        ('Clean Architecture by Robert C. Martin', 'https://blog.cleancoder.com/uncle-bob/2012/08/13/the-clean-architecture.html'),
        ('Pomodoro Technique', 'https://francescocirillo.com/pages/pomodoro-technique'),
        ('Material Design', 'https://material.io/design'),
        ('Flutter Best Practices', 'https://flutter.dev/docs/development/best-practices'),
    ]
    
    for i, (title, url) in enumerate(references, 1):
        p = doc.add_paragraph(f'{i}. {title}', style='List Number')
        p.add_run(' - ')
        add_hyperlink(p, url, url)
    
    # Save document
    output_path = 'c:\\Users\\parus\\Desktop\\ZenFlow\\hive\\ZenFlow_Documentation.docx'
    doc.save(output_path)
    print(f'Documentation generated successfully: {output_path}')

if __name__ == '__main__':
    create_zenflow_documentation()
