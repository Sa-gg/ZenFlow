#!/usr/bin/env python3
"""
Script to generate ZenFlow short documentation for printing (1-2 pages, 8.5x13 inch)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_short_documentation():
    doc = Document()
    
    # Set page size to 8.5 x 13 inches (Legal size)
    section = doc.sections[0]
    section.page_height = Inches(13)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Project Title
    title = doc.add_heading('ZENFLOW: POMODORO TIMER APPLICATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor(220, 53, 69)  # Red color
    
    # Add space
    doc.add_paragraph()
    
    # Name section
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run('Developed by: Patrick R. Sagum')
    name_run.font.size = Pt(12)
    name_run.font.bold = True
    
    program_para = doc.add_paragraph()
    program_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    program_run = program_para.add_run('Bachelor of Science in Information Technology')
    program_run.font.size = Pt(11)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run('October 2025')
    date_run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # System Description Section
    doc.add_heading('SYSTEM DESCRIPTION', 2)
    
    desc_para = doc.add_paragraph()
    desc_para.add_run(
        'ZenFlow is a productivity-focused mobile application built with Flutter that implements '
        'the Pomodoro Technique for effective time management. The application helps users break '
        'work into focused 25-minute intervals (Pomodoros) separated by short breaks, with longer '
        'breaks after every 4 completed focus sessions. The app features a clean, dark-themed '
        'interface designed to minimize distractions while maximizing productivity.'
    ).font.size = Pt(11)
    
    doc.add_paragraph()
    
    purpose_para = doc.add_paragraph()
    purpose_para.add_run('Key Features:').bold = True
    purpose_para.runs[0].font.size = Pt(11)
    
    features = [
        '• Customizable Pomodoro timer with Focus (25 min), Short Break (5 min), and Long Break (15 min)',
        '• Task management system to track work items and completed Pomodoros per task',
        '• Visual cycle progress tracking with animated indicators',
        '• Audio notifications when timer sessions complete',
        '• Local data persistence using Hive database for offline functionality',
        '• Settings panel for customizing timer durations and behavior'
    ]
    
    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Screenshots Section
    doc.add_heading('APPLICATION SCREENSHOTS', 2)
    
    screenshot_desc = doc.add_paragraph()
    screenshot_desc.add_run(
        'Below are screenshots demonstrating the main interface and key features of the ZenFlow application:'
    ).font.size = Pt(10)
    screenshot_desc.add_run().font.italic = True
    
    doc.add_paragraph()
    
    # Placeholder for screenshots
    screenshot_note = doc.add_paragraph()
    screenshot_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_note.add_run('[Screenshots to be inserted here]').font.italic = True
    screenshot_note.runs[0].font.size = Pt(10)
    screenshot_note.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Add note about screenshots
    note = doc.add_paragraph()
    note.add_run('Note: ').bold = True
    note.runs[0].font.size = Pt(9)
    note.add_run(
        'Screenshots show (1) Main timer interface, (2) Task management, (3) Settings panel, '
        'and (4) Timer completion notification.'
    ).font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Input/Output Section
    doc.add_heading('SYSTEM FUNCTIONALITY', 2)
    
    # Create a table for Input/Output
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'INPUT / USER ACTIONS'
    hdr_cells[1].text = 'OUTPUT / SYSTEM RESPONSE'
    
    # Style header
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    io_data = [
        (
            'Timer Control:\n'
            '• User selects timer type (Focus/Short Break/Long Break)\n'
            '• User clicks START button\n'
            '• User can PAUSE or RESET timer',
            
            'System Response:\n'
            '• Timer countdown begins with MM:SS format display\n'
            '• Background color changes based on timer type (Red=Focus, Yellow=Short, Blue=Long)\n'
            '• Pause button replaces Start button when active\n'
            '• Timer state persists across app sessions'
        ),
        (
            'Task Management:\n'
            '• User enters task title in input field\n'
            '• User clicks add button to create task\n'
            '• User marks task as complete with checkbox\n'
            '• User edits or deletes existing tasks',
            
            'System Response:\n'
            '• New task appears in active tasks list\n'
            '• Task moves to completed section when checked\n'
            '• Pomodoro counter increments for associated tasks\n'
            '• All tasks saved locally in Hive database\n'
            '• Edit dialog allows title modification'
        ),
        (
            'Settings Configuration:\n'
            '• User opens settings (gear icon)\n'
            '• User adjusts timer durations (Focus: 1-60 min, Breaks: 1-30 min)\n'
            '• User sets cycles before long break (2-8)\n'
            '• User toggles auto-switch option\n'
            '• User clicks APPLY or RESET ALL',
            
            'System Response:\n'
            '• Settings modal displays current values\n'
            '• Timer durations update immediately\n'
            '• New settings persist in local storage\n'
            '• Auto-switch automatically starts next timer\n'
            '• Reset ALL clears all cycle counters to zero'
        ),
        (
            'Cycle Progress:\n'
            '• User completes focus sessions\n'
            '• System tracks progress through 4-session cycle',
            
            'System Response:\n'
            '• Visual dots show cycle status:\n'
            '  - Green checkmark: Completed session\n'
            '  - Red pulsing ring: Current session\n'
            '  - Grey dot: Upcoming session\n'
            '• Completed cycles counter increments\n'
            '• Long break triggered after 4 focus sessions'
        ),
        (
            'Timer Completion:\n'
            '• Timer reaches 00:00',
            
            'System Response:\n'
            '• Audio notification plays (system alert sound)\n'
            '• Notification overlay appears with:\n'
            '  - Completion message and icon\n'
            '  - Next timer type suggestion\n'
            '• User taps anywhere to dismiss\n'
            '• Timer automatically switches to next type\n'
            '• Session data saved to history'
        ),
    ]
    
    for input_text, output_text in io_data:
        row_cells = table.add_row().cells
        
        # Input cell
        input_para = row_cells[0].paragraphs[0]
        input_para.text = input_text
        input_para.runs[0].font.size = Pt(9)
        
        # Output cell
        output_para = row_cells[1].paragraphs[0]
        output_para.text = output_text
        output_para.runs[0].font.size = Pt(9)
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(3.5)
        row.cells[1].width = Inches(4.0)
    
    doc.add_paragraph()
    
    # Technical Architecture
    doc.add_heading('TECHNICAL ARCHITECTURE', 2)
    
    tech_para = doc.add_paragraph()
    tech_para.add_run('Architecture Pattern: ').bold = True
    tech_para.runs[0].font.size = Pt(10)
    tech_para.add_run(
        'Clean Architecture with MVVM (Model-View-ViewModel) pattern, ensuring separation of '
        'concerns and maintainable code structure.'
    ).font.size = Pt(10)
    
    tech_para2 = doc.add_paragraph()
    tech_para2.add_run('Technology Stack: ').bold = True
    tech_para2.runs[0].font.size = Pt(10)
    tech_para2.add_run(
        'Flutter 3.5.3, Dart, Hive (NoSQL database), Provider (state management), '
        'Audioplayers (sound notifications), Google Fonts.'
    ).font.size = Pt(10)
    
    tech_para3 = doc.add_paragraph()
    tech_para3.add_run('Data Persistence: ').bold = True
    tech_para3.runs[0].font.size = Pt(10)
    tech_para3.add_run(
        'All data stored locally using Hive database with three main collections: Tasks, '
        'Timer Sessions, and Settings. No internet connection required.'
    ).font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Example Data Section
    doc.add_heading('EXAMPLE DATA', 2)
    
    example_para = doc.add_paragraph()
    example_para.add_run('Sample Task Data:').bold = True
    example_para.runs[0].font.size = Pt(10)
    
    example_code = '''
Task: "Complete Flutter project documentation"
- Status: Active
- Pomodoros Completed: 3
- Created: October 21, 2025

Task: "Review code architecture"
- Status: Completed
- Pomodoros Completed: 2
- Created: October 20, 2025
'''
    
    code_para = doc.add_paragraph(example_code, style='No Spacing')
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(9)
    code_para.paragraph_format.left_indent = Inches(0.25)
    
    example_para2 = doc.add_paragraph()
    example_para2.add_run('Sample Timer Session:').bold = True
    example_para2.runs[0].font.size = Pt(10)
    
    session_code = '''
Session Type: Focus
Duration: 25 minutes
Start Time: 2:30 PM
End Time: 2:55 PM
Status: Completed
Associated Task: "Complete Flutter project documentation"
'''
    
    code_para2 = doc.add_paragraph(session_code, style='No Spacing')
    code_para2.runs[0].font.name = 'Courier New'
    code_para2.runs[0].font.size = Pt(9)
    code_para2.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        '────────────────────────────────────────────────────────\n'
        'ZenFlow - Built with Flutter & Clean Architecture Principles\n'
        'Version 1.0.0 | October 2025'
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    output_path = 'c:\\Users\\parus\\Desktop\\ZenFlow\\hive\\ZenFlow_PrintDocument.docx'
    doc.save(output_path)
    print(f'Print documentation generated successfully: {output_path}')
    print('Document format: 8.5 x 13 inch (Legal size)')
    print('Note: Please insert screenshots before printing')

if __name__ == '__main__':
    create_short_documentation()
